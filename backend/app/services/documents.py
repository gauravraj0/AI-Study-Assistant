"""Document ingestion: extraction -> chunking -> embedding -> vector store."""

import io
import os
import re

from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session

from .. import models
from ..config import settings
from ..llm import get_embedder, get_provider
from ..rag import get_vector_store
from .activity import log_activity

SUPPORTED = {"pdf", "docx", "txt", "md", "csv", "json"}
TEXT_EXTS = {"txt", "md", "csv", "json"}


def _extract(data: bytes, ext: str) -> tuple[str, int]:
    if ext in TEXT_EXTS:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1", errors="replace")
        return text, max(1, text.count("\n") // 40 + 1)
    if ext == "pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise ValueError("PDF is password-protected.")
        pages = [(p.extract_text() or "") for p in reader.pages]
        return "\n\n".join(pages), len(reader.pages)
    if ext == "docx":
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        text = "\n".join(parts)
        return text, max(1, len(parts) // 30)
    raise ValueError(f"Unsupported file type: {ext}")


def chunk_text(text: str, chunk_size: int = 900, overlap: int = 120) -> list[str]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    cur = ""
    for p in paragraphs:
        if len(p) > chunk_size:
            if cur:
                chunks.append(cur)
                cur = ""
            step = chunk_size - overlap
            for i in range(0, len(p), step):
                piece = p[i : i + chunk_size]
                if len(piece) >= 60:
                    chunks.append(piece)
            continue
        if cur and len(cur) + len(p) + 2 > chunk_size:
            chunks.append(cur)
            cur = p
        else:
            cur = f"{cur}\n\n{p}" if cur else p
    if cur:
        chunks.append(cur)
    return [c for c in chunks if len(c) >= 40]


def ingest_document(db: Session, user: models.User, file: UploadFile, content: bytes) -> models.Document:
    ext = os.path.splitext(file.filename or "")[1].lower().lstrip(".")
    if ext not in SUPPORTED:
        raise HTTPException(status_code=400, detail=f"Unsupported file type '.{ext or '?'}'. Supported: {', '.join(sorted(SUPPORTED))}")
    if len(content) > settings.max_upload_bytes:
        raise HTTPException(status_code=413, detail="File too large (max 25 MB).")

    doc = models.Document(
        user_id=user.id,
        original_name=file.filename or "document",
        file_type=ext,
        status="processing",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    try:
        text, pages = _extract(content, ext)
        if len(text.strip()) < 40:
            raise ValueError("Could not extract readable text from this file.")
        chunks = chunk_text(text)
        if not chunks:
            raise ValueError("Document is too short to chunk.")

        store = get_vector_store()
        embedder = get_embedder()
        vectors = embedder.embed(chunks)
        items = [
            {"id": f"doc{doc.id}-c{i}", "text": c, "meta": {"index": i, "source": doc.original_name}}
            for i, c in enumerate(chunks)
        ]
        store.add(doc.id, items, vectors)

        summary = get_provider().summarize(text)
        doc.num_pages = pages
        doc.num_chunks = len(chunks)
        doc.status = "ready"
        doc.summary = summary
        log_activity(db, user, "document_uploaded", doc.id, f"Uploaded “{file.filename}” ({len(chunks)} chunks)")
        db.commit()
        db.refresh(doc)
        return doc
    except HTTPException:
        db.rollback()
        raise
    except Exception as e:  # noqa: BLE001
        doc.status = "error"
        doc.error = str(e)
        db.commit()
        db.refresh(doc)
        raise HTTPException(status_code=400, detail=f"Failed to process document: {e}")


def delete_document(db: Session, user: models.User, doc: models.Document) -> None:
    # null out references kept for history (quizzes/flashcards keep their results)
    for q in doc.quizzes:
        q.document_id = None
    for fs in doc.flashcard_sets:
        fs.document_id = None
    for a in db.query(models.Activity).filter(
        models.Activity.user_id == user.id, models.Activity.document_id == doc.id
    ).all():
        a.document_id = None
    get_vector_store().delete(doc.id)
    db.delete(doc)
    log_activity(db, user, "document_deleted", None, f"Deleted “{doc.original_name}”")
    db.commit()


def resummarize(db: Session, user: models.User, doc: models.Document, max_words: int = 180) -> models.Document:
    """Regenerate the AI summary for a document."""
    store = get_vector_store()
    chunks = store.all_chunks(doc.id, limit=200)
    text = "\n\n".join(c["text"] for c in chunks)
    doc.summary = get_provider().summarize(text, max_words=max_words)
    log_activity(db, user, "summary_generated", doc.id, f"Regenerated summary of “{doc.original_name}”")
    db.commit()
    db.refresh(doc)
    return doc
